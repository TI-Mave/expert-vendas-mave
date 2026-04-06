"""
Agente Comercial Mave - Servidor com RAG
=========================================
Usa ChromaDB para buscar apenas os trechos relevantes
dos documentos, economizando tokens em cada chamada.

Todos os 15 documentos ficam indexados no banco vetorial.
Cada pergunta busca os trechos mais relevantes antes de
enviar pro Claude.
"""

import os
import glob
import base64
import io
from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import HTMLResponse, JSONResponse
import anthropic
import chromadb
from chromadb.utils import embedding_functions

# ============================================================
# CONFIGURACAO
# ============================================================

MODELO = "claude-sonnet-4-20250514"
MAX_TOKENS = 4096
PASTA_BASE = os.path.dirname(os.path.abspath(__file__))

# Quantos trechos buscar por pergunta
NUM_RESULTADOS = 15

# Tamanho de cada pedaco de texto (em caracteres)
TAMANHO_CHUNK = 1500
SOBREPOSICAO = 200

TIPOS_IMAGEM = {".png", ".jpg", ".jpeg", ".gif", ".webp"}
TIPOS_TEXTO = {".txt", ".csv", ".md", ".json", ".log"}
TIPOS_PDF = {".pdf"}
TIPOS_DOCX = {".docx"}
TODOS_TIPOS = TIPOS_IMAGEM | TIPOS_TEXTO | TIPOS_PDF | TIPOS_DOCX
MAX_TAMANHO = 10 * 1024 * 1024

# ============================================================
# SYSTEM PROMPT (sempre enviado, sem os documentos)
# ============================================================

def carregar_system_prompt():
    caminho = os.path.join(PASTA_BASE, "doc___PROMPT_MANUS.txt")
    try:
        with open(caminho, "r", encoding="utf-8") as f:
            prompt = f.read()
        print(f"  System prompt carregado ({len(prompt):,} chars)")
        return prompt
    except FileNotFoundError:
        print("  AVISO: PROMPT_MANUS.txt nao encontrado")
        return "Voce e o agente comercial interno da Mave. Ajude o vendedor."

# ============================================================
# CHUNKING - dividir documentos em pedacos
# ============================================================

def dividir_em_chunks(texto, nome_doc, tamanho=TAMANHO_CHUNK, sobreposicao=SOBREPOSICAO):
    chunks = []
    inicio = 0
    while inicio < len(texto):
        fim = inicio + tamanho

        if fim < len(texto):
            pos_paragrafo = texto.rfind("\n\n", inicio + tamanho // 2, fim + 200)
            if pos_paragrafo > inicio:
                fim = pos_paragrafo
            else:
                pos_frase = texto.rfind(". ", inicio + tamanho // 2, fim + 100)
                if pos_frase > inicio:
                    fim = pos_frase + 1

        pedaco = texto[inicio:fim].strip()
        if len(pedaco) > 50:
            chunks.append({
                "texto": pedaco,
                "documento": nome_doc,
                "posicao": len(chunks),
            })

        inicio = fim - sobreposicao
        if inicio < 0:
            inicio = 0
        if fim >= len(texto):
            break

    return chunks

# ============================================================
# INDEXAR DOCUMENTOS NO CHROMADB
# ============================================================

def criar_indice():
    print("\n" + "=" * 50)
    print("Indexando documentos no ChromaDB...")
    print("=" * 50)

    ef = embedding_functions.DefaultEmbeddingFunction()
    cliente_chroma = chromadb.Client()

    try:
        cliente_chroma.delete_collection("mave_docs")
    except Exception:
        pass

    colecao = cliente_chroma.create_collection(
        name="mave_docs",
        embedding_function=ef,
        metadata={"hnsw:space": "cosine"},
    )

    arquivos = sorted(glob.glob(os.path.join(PASTA_BASE, "doc___*.txt")))
    total_chunks = 0

    for caminho in arquivos:
        nome_arquivo = os.path.basename(caminho)
        nome_limpo = nome_arquivo.replace("doc___", "")

        if nome_limpo == "PROMPT_MANUS.txt":
            continue

        try:
            with open(caminho, "r", encoding="utf-8") as f:
                conteudo = f.read()

            chunks = dividir_em_chunks(conteudo, nome_limpo)

            if chunks:
                ids = [f"{nome_limpo}_{i}" for i in range(len(chunks))]
                textos = [c["texto"] for c in chunks]
                metadados = [{"documento": c["documento"], "posicao": c["posicao"]} for c in chunks]

                colecao.add(
                    ids=ids,
                    documents=textos,
                    metadatas=metadados,
                )

                total_chunks += len(chunks)
                print(f"  {nome_limpo}: {len(chunks)} chunks")

        except Exception as e:
            print(f"  ERRO ao indexar {nome_limpo}: {e}")

    print(f"\nTotal: {total_chunks} chunks indexados")
    print("=" * 50)

    return colecao


SYSTEM_PROMPT = carregar_system_prompt()
COLECAO = criar_indice()

# ============================================================
# BUSCA RAG
# ============================================================

def buscar_contexto(pergunta, n_resultados=NUM_RESULTADOS):
    try:
        resultados = COLECAO.query(
            query_texts=[pergunta],
            n_results=n_resultados,
        )

        trechos = []
        docs_usados = set()

        if resultados and resultados["documents"] and resultados["documents"][0]:
            for i, texto in enumerate(resultados["documents"][0]):
                meta = resultados["metadatas"][0][i] if resultados["metadatas"] else {}
                doc_nome = meta.get("documento", "desconhecido")
                docs_usados.add(doc_nome)
                trechos.append(f"[Fonte: {doc_nome}]\n{texto}")

        contexto = "\n\n---\n\n".join(trechos)
        print(f"  RAG: {len(trechos)} trechos de {len(docs_usados)} docs ({len(contexto):,} chars)")
        return contexto

    except Exception as e:
        print(f"  Erro no RAG: {e}")
        return ""

# ============================================================
# PROCESSAMENTO DE ARQUIVOS DO VENDEDOR
# ============================================================

def extrair_texto_pdf(conteudo_bytes):
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(conteudo_bytes))
        textos = [p.extract_text() for p in reader.pages if p.extract_text()]
        return "\n\n".join(textos) if textos else "[PDF sem texto extraivel]"
    except Exception as e:
        return f"[Erro ao ler PDF: {e}]"


def extrair_texto_docx(conteudo_bytes):
    try:
        from docx import Document
        doc = Document(io.BytesIO(conteudo_bytes))
        textos = [p.text for p in doc.paragraphs if p.text.strip()]
        for tabela in doc.tables:
            for linha in tabela.rows:
                celulas = [c.text.strip() for c in linha.cells if c.text.strip()]
                if celulas:
                    textos.append(" | ".join(celulas))
        return "\n".join(textos) if textos else "[DOCX sem texto]"
    except Exception as e:
        return f"[Erro ao ler DOCX: {e}]"


async def processar_arquivo(arquivo: UploadFile):
    nome = arquivo.filename or "arquivo"
    ext = os.path.splitext(nome)[1].lower()

    if ext not in TODOS_TIPOS:
        return ("texto", f"[Tipo nao suportado: {ext}. Envie PDF, imagem, .txt, .csv ou .docx]")

    conteudo = await arquivo.read()
    if len(conteudo) > MAX_TAMANHO:
        return ("texto", f"[Arquivo muito grande: {len(conteudo)/(1024*1024):.1f}MB. Maximo: 10MB]")

    if ext in TIPOS_IMAGEM:
        media_types = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".gif": "image/gif", ".webp": "image/webp"}
        return ("imagem", {
            "tipo_media": media_types[ext],
            "dados_b64": base64.b64encode(conteudo).decode("utf-8"),
            "nome": nome,
        })

    if ext in TIPOS_TEXTO:
        try:
            texto = conteudo.decode("utf-8")
        except UnicodeDecodeError:
            texto = conteudo.decode("latin-1")
        return ("texto", f"[Arquivo '{nome}']\n{texto}")

    if ext in TIPOS_PDF:
        return ("texto", f"[PDF '{nome}']\n{extrair_texto_pdf(conteudo)}")

    if ext in TIPOS_DOCX:
        return ("texto", f"[DOCX '{nome}']\n{extrair_texto_docx(conteudo)}")

    return ("texto", f"[Nao processado: {nome}]")

# ============================================================
# SERVIDOR
# ============================================================

app = FastAPI(title="Agente Comercial Mave")
cliente_anthropic = anthropic.Anthropic()
conversas = {}


@app.get("/", response_class=HTMLResponse)
async def pagina_principal():
    caminho = os.path.join(PASTA_BASE, "index.html")
    with open(caminho, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.post("/chat")
async def chat(
    mensagem: str = Form(default=""),
    session_id: str = Form(default="default"),
    arquivo: UploadFile = File(default=None),
):
    try:
        mensagem = mensagem.strip()
        if not mensagem and (not arquivo or not arquivo.filename):
            return JSONResponse(content={"erro": "Envie uma mensagem ou arquivo"}, status_code=400)

        if session_id not in conversas:
            conversas[session_id] = []
        historico = conversas[session_id]

        content_parts = []
        nome_arquivo = None
        texto_arquivo = ""

        if arquivo and arquivo.filename:
            nome_arquivo = arquivo.filename
            tipo, dados = await processar_arquivo(arquivo)

            if tipo == "imagem":
                content_parts.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": dados["tipo_media"],
                        "data": dados["dados_b64"],
                    }
                })
                if not mensagem:
                    mensagem = f"O vendedor enviou a imagem '{dados['nome']}'. Analise e ajude conforme o contexto comercial."
            else:
                texto_arquivo = dados
                if not mensagem:
                    mensagem = f"O vendedor enviou o arquivo '{nome_arquivo}'. Analise o conteudo e ajude."

        # BUSCA RAG
        busca_texto = mensagem
        if texto_arquivo:
            busca_texto = mensagem + " " + texto_arquivo[:500]
        contexto_rag = buscar_contexto(busca_texto)

        # Monta system prompt com contexto relevante
        system_com_contexto = f"""{SYSTEM_PROMPT}

<contexto_relevante>
Os trechos abaixo foram selecionados da base de conhecimento da Mave como os mais relevantes para esta pergunta. Use-os para fundamentar sua resposta.

{contexto_rag}
</contexto_relevante>
"""

        # Monta mensagem do usuario
        if texto_arquivo:
            content_parts.append({"type": "text", "text": texto_arquivo})
        content_parts.append({"type": "text", "text": mensagem})
        historico.append({"role": "user", "content": content_parts})

        if len(historico) > 20:
            historico = historico[-20:]
            conversas[session_id] = historico

        resposta = cliente_anthropic.messages.create(
            model=MODELO,
            max_tokens=MAX_TOKENS,
            system=system_com_contexto,
            messages=historico,
        )

        texto_resposta = resposta.content[0].text
        historico.append({"role": "assistant", "content": texto_resposta})

        return JSONResponse(content={"resposta": texto_resposta, "arquivo_recebido": nome_arquivo})

    except anthropic.APIError as e:
        print(f"Erro API: {e}")
        return JSONResponse(content={"erro": f"Erro na API: {str(e)}"}, status_code=500)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(content={"erro": f"Erro interno: {str(e)}"}, status_code=500)


@app.post("/nova-conversa")
async def nova_conversa(request: Request):
    dados = await request.json()
    session_id = dados.get("session_id", "default")
    conversas[session_id] = []
    return JSONResponse(content={"status": "ok"})


@app.get("/saude")
async def saude():
    n_chunks = COLECAO.count() if COLECAO else 0
    return {"status": "ok", "chunks_indexados": n_chunks}
